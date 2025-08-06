from flask import session
from utility.rag_processing import process_pdf_for_rag
from utility.converting_text_to_audio import convert_text_to_audio
from utility.gemini_summarize_tool import gemini_summarize
import os
import time
import docx
import uuid
from pathlib import Path
import asyncio
from dotenv import load_dotenv
import google.generativeai as genai
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain.schema import Document
from langchain_community.vectorstores import FAISS


load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
if api_key:
    genai.configure(api_key=api_key)
else:
    print("GEMINI_API_KEY not found in .env file. Some features might not work.")


BASE_DIR = Path(__file__).parent.parent

UPLOAD_FOLDER = BASE_DIR / "uploads"
STATIC_FOLDER = BASE_DIR / "static"
AUDIO_FOLDER = STATIC_FOLDER / "audio"
IMAGES_FOLDER = STATIC_FOLDER / "images"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(AUDIO_FOLDER, exist_ok=True)
os.makedirs(IMAGES_FOLDER, exist_ok=True)


def process_uploaded_file(file):
    """
    Saves the uploaded file and processes it to extract text and images.
    """
    # Remove the blanket file deletion. Cleanup will now be handled by clean_up route.
    # for folder in [UPLOAD_FOLDER, IMAGES_FOLDER]:
    #     for filename in os.listdir(folder):
    #         file_path = os.path.join(folder, filename)
    #         if os.path.isfile(file_path):
    #             os.remove(file_path)

    filepath = os.path.join(UPLOAD_FOLDER, f"{session.get('user_id')}_" + file.filename)
    file.save(filepath)

    file_extension = os.path.splitext(file.filename.lower())[1]
    file_type = file_extension[1:]

    text_chunks, image_info, full_text = [], [], ""

    if file_type == 'pdf':
        text_chunks, image_info, full_text = process_pdf_for_rag(filepath, str(BASE_DIR))
    elif file_type in ['doc', 'docx']:
        doc = docx.Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        for i, para_text in enumerate(paragraphs):
            text_chunks.append({"text": para_text, "page": i + 1})

    return text_chunks, image_info, full_text, file_type, filepath

def summarize_text(full_text, text_chunks, image_info):
    if not api_key:
        return "GEMINI_API_KEY not found."

    summary_start_time = time.time()
    summary = gemini_summarize(full_text)
    summary_end_time = time.time()
    count = len(summary.split())
    print(f"Summary generation time: {int(summary_end_time - summary_start_time)} seconds And Count: {count}")

    if not text_chunks:
        return [{"response": summary, "page": "N/A", "source_text": "Could not extract source text.", "images": []}]

    original_docs = [Document(p["text"], metadata={"page": p["page"]}) for p in text_chunks]

    try:
        asyncio.get_running_loop()
    except RuntimeError:
        asyncio.set_event_loop(asyncio.new_event_loop())

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001", google_api_key=api_key
    )
    original_store = FAISS.from_documents(original_docs, embeddings)

    final_response_with_citations = []
    summarized_pages = set()

    for para in [p.strip() for p in summary.split("\n\n") if p.strip()]:
        try:
            match = original_store.similarity_search(para, k=1)[0]
            page_num = int(match.metadata["page"])
            
            page_images = [img["path"] for img in image_info if img["page"] == page_num]

            if page_num not in summarized_pages:
                final_response_with_citations.append(
                    {
                        "response": para,
                        "page": str(page_num),
                        "source_text": match.page_content,
                        "images": page_images,
                    }
                )
                summarized_pages.add(page_num)
            else:
                # Find the existing entry and append the text
                for item in final_response_with_citations:
                    if item["page"] == str(page_num):
                        item["response"] += "\n\n" + para
                        break

        except Exception as e:
            print(f"Similarity search failed for a paragraph: {e}")
            final_response_with_citations.append(
                {
                    "response": para,
                    "page": "not found",
                    "source_text": "Source text not found",
                    "images": [],
                }
            )
    
    return final_response_with_citations

def convert_summary_to_audio(summary):
    audio_start_time = time.time()
    audio_filename = f"{session.get('user_id')}_audio.wav"
    audio_file_path = os.path.join(AUDIO_FOLDER, audio_filename)
    
    if isinstance(summary, list):
        summary_text = "\n".join([item['response'] for item in summary])
    else:
        summary_text = summary

    convert_text_to_audio(summary_text, audio_file_path)
    audio_end_time = time.time()
    print(f"Audio generation time: {int(audio_end_time - audio_start_time)} seconds")
    return audio_filename

